# import os
# import pdfplumber
# # import fitz  # PyMuPDF
# from langchain_community.document_loaders.pdf import PyPDFLoader
# from docx import Document as DocxDocument
# from fastapi import FastAPI, HTTPException, Form
# from fastapi.responses import HTMLResponse
# from pydantic import BaseModel
# from google.oauth2 import service_account
# from googleapiclient.discovery import build
# from googleapiclient.http import MediaIoBaseDownload
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.document_loaders import TextLoader
# from langchain_mistralai.embeddings import MistralAIEmbeddings
# from langchain_community.vectorstores import FAISS
# from langchain.chains import create_retrieval_chain
# from langchain.prompts import ChatPromptTemplate
# from langchain.chains.combine_documents import create_stuff_documents_chain
# from langchain_mistralai.chat_models import ChatMistralAI
# import io
# import json
# from typing import List
# from langchain.docstore.document import Document
# from mistralai.client import MistralClient
# from mistralai import Mistral
# #from mistralai.models.chat_completion import ChatMessage
# import sqlite3

# app = FastAPI()

# class Query(BaseModel):
#     query: str

# class QueryRequest(BaseModel):
#     query: str

# folder_id = None
# client = Mistral(api_key="")

# # Global variables to cache documents and vector store
# cached_documents = []
# cached_vector_store = None

# # Path to cache file
# CACHE_FILE = 'document_cache.json'




# DATABASE_FILE = 'history.db'

# # Initialize SQLite database
# def init_db():
#     conn = sqlite3.connect(DATABASE_FILE)
#     cursor = conn.cursor()
#     cursor.execute('''CREATE TABLE IF NOT EXISTS history (
#                         id INTEGER PRIMARY KEY AUTOINCREMENT,
#                         query TEXT,
#                         answer TEXT
#                     )''')
#     conn.commit()
#     conn.close()

# init_db()

# # Save history to the database
# def save_history(query: str, answer: str):
#     conn = sqlite3.connect(DATABASE_FILE)
#     cursor = conn.cursor()
#     cursor.execute('INSERT INTO history (query, answer) VALUES (?, ?)', (query, answer))
#     conn.commit()
#     conn.close()

# # Fetch all history
# @app.get("/history")
# async def get_history():
#     conn = sqlite3.connect(DATABASE_FILE)
#     cursor = conn.cursor()
#     cursor.execute('SELECT query, answer FROM history')
#     rows = cursor.fetchall()
#     conn.close()
#     return [{'query': row[0], 'answer': row[1]} for row in rows]

# # Clear all history
# @app.delete("/history")
# async def clear_history():
#     conn = sqlite3.connect(DATABASE_FILE)
#     cursor = conn.cursor()
#     cursor.execute('DELETE FROM history')
#     conn.commit()
#     conn.close()
#     return {'message': 'History cleared successfully'}



import os
import pdfplumber
from langchain_community.document_loaders.pdf import PyPDFLoader
from docx import Document as DocxDocument
from fastapi import FastAPI, HTTPException, Form, Depends, Request
from fastapi.responses import HTMLResponse, RedirectResponse
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from pydantic import BaseModel
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.chains import create_retrieval_chain
from langchain.prompts import ChatPromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_mistralai.embeddings import MistralAIEmbeddings
from langchain_mistralai.chat_models import ChatMistralAI
import io
import json
from typing import List
from langchain.docstore.document import Document
from mistralai import Mistral
from pymongo import MongoClient
from passlib.context import CryptContext
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import os
from dotenv import load_dotenv
from datetime import datetime

load_dotenv()
print("MONGO_URI:", os.getenv("MONGO_URI")) 
app = FastAPI()
class Query(BaseModel):
    query: str

class QueryRequest(BaseModel):
    query: str 

app = FastAPI()


# added externally
CACHE_FILE = 'document_cache.json'

# Setup MongoDB connection
MONGO_URI = os.getenv("MONGO_URI")
client = MongoClient(MONGO_URI)
db = client["user_db"]
users_collection = db["users"]
history_collection = db["history"]

# Password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/login")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# Authentication Functions
def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def authenticate_user(username: str, password: str):
    user = users_collection.find_one({"username": username})
    if user and verify_password(password, user["password"]):
        return user
    return None

@app.post("/register")
async def register(username: str = Form(...), password: str = Form(...)):
    if users_collection.find_one({"username": username}):
        raise HTTPException(status_code=400, detail="Username already exists")
    
    hashed_password = get_password_hash(password)
    result = users_collection.insert_one({"username": username, "password": hashed_password})
    
    if result.inserted_id:
        print(f"User registered with ID: {result.inserted_id}")
    else:
        print("Failed to insert user into the database")
    response = RedirectResponse(url="/", status_code=302)
    response.set_cookie(key="access_token", value=username, httponly=True)  # Secure the cookie
    return response

# @app.post("/register")
# async def register(username: str = Form(...), password: str = Form(...)):
#     if users_collection.find_one({"username": username}):
#         raise HTTPException(status_code=400, detail="Username already exists")
#     hashed_password = get_password_hash(password)
#     users_collection.insert_one({"username": username, "password": hashed_password})
#     return {"message": "User registered successfully"}

@app.post("/login")
async def login(form_data: OAuth2PasswordRequestForm = Depends()):
    user = authenticate_user(form_data.username, form_data.password)
    if not user:
        raise HTTPException(status_code=400, detail="Invalid credentials")
    response = RedirectResponse(url="/", status_code=302)
    response.set_cookie(key="access_token", value=form_data.username)
    return response


@app.get("/logout")
async def logout():
    print("Logout route accessed")
    response = RedirectResponse(url="/login_page", status_code=302)
    response.delete_cookie(key="access_token", path="/")
    response.set_cookie(key="access_token", value="", expires=0)
    response.headers["Cache-Control"] = "no-store"
    return response

# @app.get("/logout")
# async def logout():
#     response = RedirectResponse(url="/login_page", status_code=302)
#     response.delete_cookie(key="access_token")
#     return response

@app.get("/login_page", response_class=HTMLResponse)
async def login_page():
    with open("app.html", "r") as file:
        return HTMLResponse(content=file.read(), status_code=200)
    
@app.get("/register_page", response_class=HTMLResponse)
async def register_page():
    with open("register.html", "r") as file:
        return HTMLResponse(content=file.read(), status_code=200)

# Middleware to check authentication
async def get_current_user(request: Request):
    token = request.cookies.get("access_token")
    if not token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    return token

# Save query and answer to MongoDB
def save_chat(user_id, username, role, message):
    """
    Save a chat message in MongoDB under a specific user.
    """
    chat_entry = {
        "timestamp": datetime.utcnow().isoformat(),
        "role": role,  # 'user' or 'bot'
        "message": message
    }

    # Update the user's chat history
    history_collection.update_one(
        {"user_id": user_id},  # Find by user_id
        {"$push": {"chats": chat_entry}, "$set": {"username": username}},  # Append chat
        upsert=True  # Create if not exists
    )

# def save_history(query: str, answer: str):
#     history_collection.insert_one({"query": query, "answer": answer})

# Fetch all history
# @app.get("/history")
# async def get_history(user: str = Depends(get_current_user)):
#     history = list(history_collection.find({}, {"_id": 0}))
#     return history


@app.get("/history")
async def get_user_history(user: str = Depends(get_current_user)):
    user_data = history_collection.find({"user": user})

    queries = list(user_data)

    if not queries:
        return {"message": "No queries found for this user."}

    history_by_date = {}

    for query in queries:
        timestamp = query.get("_id").generation_time  # Get the timestamp from ObjectId

        # Extract date from ObjectId timestamp
        date = timestamp.strftime("%Y-%m-%d")

        # Group by date
        if date not in history_by_date:
            history_by_date[date] = []

        # Append query/answer pairs
        history_by_date[date].append({
            "query": query.get("query"),
            "answer": query.get("answer")
        })

    return {"history": history_by_date}

@app.get("/history_page", response_class=HTMLResponse)
async def history_page():
    with open("history.html", "r") as file:
        return HTMLResponse(content=file.read(), status_code=200)


# @app.get("/history")
# async def get_user_history(user: str = Depends(get_current_user)):
#     user_history = list(history_collection.find({"user": user}, {"_id": 0, "query": 1, "answer": 1, "timestamp": 1}).sort("timestamp", -1))
    
#     history_by_date = {}

#     # Group queries by date
#         # Group queries by date
#     for entry in user_history:
#         # Handle different timestamp formats
#         timestamp = entry.get("timestamp")
#         if isinstance(timestamp, str):
#             try:
#                 date = datetime.fromisoformat(timestamp).strftime("%Y-%m-%d")
#             except ValueError:
#                 date = timestamp.split("T")[0]  # Fallback for timestamps like '2024-10-27T12:34:56'
#         elif isinstance(timestamp, datetime):
#             date = timestamp.strftime("%Y-%m-%d")
#         else:
#             date = "Unknown Date"

#         # Group by date
#         if date not in history_by_date:
#             history_by_date[date] = []

#         history_by_date[date].append({
#             "query": entry["query"],
#             "answer": entry["answer"]
#         })

#     if not user_history:
#         return {"message": "No queries found for this user."}

    # Convert ObjectId to string and return the history
    # for record in user_history:
    #     record["_id"] = str(record["_id"])
    # history=list(user_history)
    
    # return {"history": history_by_date}
# @app.get("/history")
# async def get_history(user: str = Depends(get_current_user)):
#     history = list(history_collection.find({"user": user}, {"_id": 0}))
#     return history

# @app.delete("/history")
# async def clear_history(user: str = Depends(get_current_user)):
#     history_collection.delete_many({})
#     return {'message': 'History cleared successfully'}



# @app.post("/query")
# async def handle_query(query: str = Form(...)):
#     # Your query processing logic here
#     print(f"Query received: {query}")
#     return RedirectResponse(url="/history", status_code=303)
# @app.post("/query")
# async def query_api(request: QueryRequest):
#     return {"response": f"You asked: {request.query}"}
# @app.post("/query")
# async def query(request: Request):
#     data = await request.json()
#     print("Received Data:", data)  # Debugging
#     return {"message": "Success"}


@app.get("/", response_class=HTMLResponse)
async def read_root(request: Request, user: str = Depends(get_current_user)):
    if not user:
        return RedirectResponse(url="/login_page", status_code=303)

    with open("index.html", "r") as file:
        return HTMLResponse(content=file.read(), status_code=200)
    
# @app.get("/", response_class=HTMLResponse)
# async def read_root(user: str = Depends(get_current_user)):
#     with open("index.html", "r") as file:
#         return HTMLResponse(content=file.read(), status_code=200)


def save_history(query: str, answer: str, user: str):
    history_collection.insert_one({"query": query, "answer": answer, "user": user})


def save_cache(documents):
    with open(CACHE_FILE, 'w') as f:
        json.dump([doc.dict() for doc in documents], f)

def load_cache():
    if os.path.exists(CACHE_FILE):
        with open(CACHE_FILE, 'r') as f:
            return [Document(**doc) for doc in json.load(f)]
    return []

@app.post("/set_folder/")
async def set_folder(folder_id_input: str = Form(...)):
    global folder_id, cached_documents, cached_vector_store
    folder_id = folder_id_input

    SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
    SERVICE_ACCOUNT_FILE = r'C:\Users\Pawan\Downloads\AI-Search-Engine\drive-with-rag-bf7e8c4b530b.json'


    credentials = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    service = build('drive', 'v3', credentials=credentials)

    def list_files_in_folder(service, folder_id):
        query = f"'{folder_id}' in parents and trashed=false"
        try:
            results = service.files().list(q=query, pageSize=10, fields="files(id, name)").execute()
            items = results.get('files', [])
            if not items:
                print(f"No files found in the folder with ID: {folder_id}")
            return items
        except Exception as e:
            print(f"An error occurred: {e}")
            return []

    files = list_files_in_folder(service, folder_id)
    if not files:
        raise ValueError(f"No files found or unable to access the folder with ID: {folder_id}")

    documents = []
    for file in files:
        file_id = file['id']
        file_name = file['name']
        print(f"Processing file: {file_name} (ID: {file_id})")
        request = service.files().get_media(fileId=file_id)
        file_io = io.BytesIO()
        downloader = MediaIoBaseDownload(file_io, request)
        done = False
        while done is False:
            status, done = downloader.next_chunk()
        file_io.seek(0)

        if file_name.endswith('.pdf'):
            text = extract_text_from_pdf(file_io)
        elif file_name.endswith('.docx'):
            text = extract_text_from_docx(file_io)
        else:
            text = file_io.read().decode('utf-8')

        documents.append(Document(page_content=text))

    # Split text into chunks
    text_splitter = RecursiveCharacterTextSplitter()
    split_documents = text_splitter.split_documents(documents)

    os.environ["HF_TOKEN"] = "" 

    mistral_api_key = "your key"
    embeddings = MistralAIEmbeddings(model="mistral-embed", api_key=mistral_api_key)

    if not split_documents:
        raise HTTPException(status_code=500, detail="No documents found to process")

    cached_vector_store = FAISS.from_documents(split_documents, embeddings)
    cached_documents = split_documents

    # Save the cache
    save_cache(cached_documents)

    return {"message": "Folder ID set and documents indexed successfully"}

# def extract_text_from_pdf(file_io):
#     text = ""
#     document = fitz.open(stream=file_io, filetype="pdf")
#     for page_num in range(len(document)):
#         page = document.load_page(page_num)
#         text += page.get_text()
#     return text


def extract_text_from_pdf(file_io):
    text = ""
    # Open the PDF file using pdfplumber
    with pdfplumber.open(file_io) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""  # Add extracted text from each page
    return text


def extract_text_from_docx(file_io):
    doc = DocxDocument(file_io)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])


@app.post("/query")
async def query_drive(query_request: QueryRequest, user: str = Depends(get_current_user)):
    global cached_vector_store
    if not folder_id:
        raise HTTPException(status_code=400, detail="Folder ID not set")

    if not cached_documents:
        raise HTTPException(status_code=500, detail="Documents are not loaded. Please set the folder ID again.")

    retriever = cached_vector_store.as_retriever()

    mistral_api_key = "your key"
    model = ChatMistralAI(api_key=mistral_api_key)

    prompt = ChatPromptTemplate.from_template("""Answer the following question based only on the provided context:

    <context>
    {context}
    </context>

    Question: {input}""")

    document_chain = create_stuff_documents_chain(model, prompt)
    retrieval_chain = create_retrieval_chain(retriever, document_chain)

    query = query_request.query
    try:
        response = retrieval_chain.invoke({"input": query})
        answer = response["answer"]
        
        # Save query and response to MongoDB
        save_history(query=query, answer=answer, user=user)
        
        return {"query": query, "answer": answer}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
# @app.post("/query")
# async def query_drive(query_request: QueryRequest):
#     global cached_vector_store
#     if not folder_id:
#         raise HTTPException(status_code=400, detail="Folder ID not set")

#     if not cached_documents:
#         raise HTTPException(status_code=500, detail="Documents are not loaded. Please set the folder ID again.")
    # retriever = cached_vector_store.as_retriever()

    # mistral_api_key = "8larRmMd24jOubEOKrcysrCBYLjplJRn"
    # model = ChatMistralAI(api_key=mistral_api_key)

    # prompt = ChatPromptTemplate.from_template("""Answer the following question based only on the provided context:

    # <context>
    # {context}
    # </context>

    # Question: {input}""")

    # document_chain = create_stuff_documents_chain(model, prompt)
    # retrieval_chain = create_retrieval_chain(retriever, document_chain)

    # try:
    #     response = retrieval_chain.invoke({"input": query_request.query})
    #     return {"answer": response["answer"]}
    # except Exception as e:
    #     raise HTTPException(status_code=500, detail=str(e))


@app.get("/", response_class=HTMLResponse)
async def read_root():
    with open("index.html", "r") as file:
        return HTMLResponse(content=file.read(), status_code=200)
    
if __name__ == "__main__":
    import threading
    import uvicorn

    def run_app():
        uvicorn.run(app, host="0.0.0.0", port=8000)

    thread = threading.Thread(target=run_app, args=(), daemon=True)
    thread.start()
